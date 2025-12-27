from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
import os
import numpy as np
from sentence_transformers import SentenceTransformer
import gensim.downloader as api
from gensim.models.doc2vec import Doc2Vec, TaggedDocument
import PyPDF2
import docx
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import json
import re
from werkzeug.utils import secure_filename

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

try:
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('averaged_perceptron_tagger')

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize AI Models
print("🔄 Loading AI models... This may take a few minutes on first run.")
print("📦 Loading SBERT model...")
sbert_model = SentenceTransformer('all-MiniLM-L6-v2')
print("✅ SBERT loaded successfully!")

print("📦 Loading GloVe model...")
try:
    glove_model = api.load('glove-wiki-gigaword-100')
    print("✅ GloVe loaded successfully!")
except:
    print("⚠️ GloVe model will be downloaded on first use")
    glove_model = None

print("🎉 All models loaded successfully!")

# NLP Preprocessing
def preprocess_text(text):
    """Clean and preprocess text"""
    if not text:
        return [], ""
    
    # Convert to lowercase
    text = text.lower()
    # Remove special characters and digits but keep spaces
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    # Tokenize
    tokens = word_tokenize(text)
    # Remove stopwords
    stop_words = set(stopwords.words('english'))
    tokens = [w for w in tokens if w not in stop_words and len(w) > 2]
    return tokens, ' '.join(tokens)

# Text Extraction Functions
def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ''
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + '\n'
        return text
    except Exception as e:
        print(f"❌ Error extracting PDF: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        print(f"❌ Error extracting DOCX: {e}")
        return ""

def extract_text(file_path):
    """Extract text based on file extension"""
    ext = file_path.split('.')[-1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    return ""

# Similarity Calculation Functions
def calculate_sbert_similarity(text1, text2):
    """Calculate similarity using SBERT"""
    try:
        embeddings = sbert_model.encode([text1, text2])
        similarity = np.dot(embeddings[0], embeddings[1]) / (
            np.linalg.norm(embeddings[0]) * np.linalg.norm(embeddings[1])
        )
        return float(similarity)
    except Exception as e:
        print(f"❌ SBERT error: {e}")
        return 0.0

def calculate_glove_similarity(tokens1, tokens2):
    """Calculate similarity using GloVe"""
    global glove_model
    
    try:
        if glove_model is None:
            print("📦 Downloading GloVe model (one-time download)...")
            glove_model = api.load('glove-wiki-gigaword-100')
        
        def get_avg_embedding(tokens):
            embeddings = []
            for token in tokens:
                try:
                    embeddings.append(glove_model[token])
                except KeyError:
                    continue
            return np.mean(embeddings, axis=0) if embeddings else np.zeros(100)
        
        emb1 = get_avg_embedding(tokens1)
        emb2 = get_avg_embedding(tokens2)
        
        if np.linalg.norm(emb1) == 0 or np.linalg.norm(emb2) == 0:
            return 0.0
        
        similarity = np.dot(emb1, emb2) / (np.linalg.norm(emb1) * np.linalg.norm(emb2))
        return float(similarity)
    except Exception as e:
        print(f"❌ GloVe error: {e}")
        return 0.0

def calculate_doc2vec_similarity(tokens1, tokens2):
    """Calculate similarity using Doc2Vec"""
    try:
        if not tokens1 or not tokens2:
            return 0.0
        
        # Create tagged documents
        documents = [
            TaggedDocument(tokens1, [0]),
            TaggedDocument(tokens2, [1])
        ]
        
        # Train Doc2Vec model
        model = Doc2Vec(documents, vector_size=100, window=5, min_count=1, workers=4, epochs=40)
        
        # Get document vectors
        vec1 = model.dv[0]
        vec2 = model.dv[1]
        
        # Calculate cosine similarity
        similarity = np.dot(vec1, vec2) / (np.linalg.norm(vec1) * np.linalg.norm(vec2))
        return float(similarity)
    except Exception as e:
        print(f"❌ Doc2Vec error: {e}")
        return 0.0

# Mock Gemini Recommendations (replace with actual API call)
def get_gemini_recommendations(resume_text, job_description, similarity_scores):
    """Generate recommendations (Mock version - replace with actual Gemini API)"""
    # Extract keywords from job description
    jd_tokens = job_description.lower().split()
    resume_tokens = resume_text.lower().split()
    
    # Common technical skills to look for
    common_skills = ['python', 'java', 'javascript', 'react', 'node', 'sql', 'docker', 
                     'kubernetes', 'aws', 'machine learning', 'deep learning', 'nlp',
                     'flask', 'django', 'tensorflow', 'pytorch', 'git']
    
    matched_skills = []
    for skill in common_skills:
        if skill in jd_tokens and skill in resume_tokens:
            matched_skills.append(skill.title())
    
    # Generate recommendations based on score
    avg_score = similarity_scores.get('average', 0)
    
    if avg_score > 0.75:
        recommendations = [
            "Strong technical background aligned with job requirements",
            "Relevant experience demonstrated in resume",
            "Highly recommended for interview"
        ]
        assessment = "Excellent candidate - Strong match"
    elif avg_score > 0.65:
        recommendations = [
            "Good technical foundation with relevant skills",
            "Some experience gaps but trainable",
            "Recommended for screening interview"
        ]
        assessment = "Good candidate - Worth considering"
    else:
        recommendations = [
            "Basic qualifications present",
            "May require significant training",
            "Consider for entry-level roles"
        ]
        assessment = "Fair candidate - Review carefully"
    
    return {
        'matched_skills': matched_skills[:5] if matched_skills else ['General skills match'],
        'recommendations': recommendations,
        'assessment': assessment
    }

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/analyze', methods=['POST'])
def analyze_resumes():
    """Main endpoint for resume analysis"""
    try:
        print("\n" + "="*50)
        print("🚀 Starting Resume Analysis")
        print("="*50)
        
        # Get job description
        job_description = request.form.get('job_description', '')
        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400
        
        print(f"📄 Job Description length: {len(job_description)} characters")
        
        # Process job description
        jd_tokens, jd_clean = preprocess_text(job_description)
        print(f"🔤 Job Description tokens: {len(jd_tokens)}")
        
        # Get uploaded files
        files = request.files.getlist('resumes')
        if not files or files[0].filename == '':
            return jsonify({'error': 'No resumes uploaded'}), 400
        
        print(f"📁 Number of resumes uploaded: {len(files)}")
        
        results = []
        
        for idx, file in enumerate(files, 1):
            if file and allowed_file(file.filename):
                print(f"\n--- Processing Resume {idx}/{len(files)}: {file.filename} ---")
                
                # Save file
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                print(f"💾 File saved temporarily")
                
                # Extract text
                print(f"📖 Extracting text...")
                resume_text = extract_text(filepath)
                print(f"✅ Extracted {len(resume_text)} characters")
                
                if not resume_text.strip():
                    print(f"⚠️ Warning: No text extracted from {filename}")
                    os.remove(filepath)
                    continue
                
                # Preprocess
                print(f"🔄 Preprocessing text...")
                resume_tokens, resume_clean = preprocess_text(resume_text)
                print(f"✅ Generated {len(resume_tokens)} tokens")
                
                # Calculate similarities
                print(f"🧮 Calculating SBERT similarity...")
                sbert_score = calculate_sbert_similarity(jd_clean, resume_clean)
                print(f"   SBERT: {sbert_score:.4f}")
                
                print(f"🧮 Calculating GloVe similarity...")
                glove_score = calculate_glove_similarity(jd_tokens, resume_tokens)
                print(f"   GloVe: {glove_score:.4f}")
                
                print(f"🧮 Calculating Doc2Vec similarity...")
                doc2vec_score = calculate_doc2vec_similarity(jd_tokens, resume_tokens)
                print(f"   Doc2Vec: {doc2vec_score:.4f}")
                
                avg_score = (sbert_score + glove_score + doc2vec_score) / 3
                print(f"📊 Average Score: {avg_score:.4f}")
                
                # Get AI recommendations
                print(f"🤖 Generating AI recommendations...")
                similarity_scores = {
                    'sbert': sbert_score,
                    'glove': glove_score,
                    'doc2vec': doc2vec_score,
                    'average': avg_score
                }
                gemini_insights = get_gemini_recommendations(
                    resume_text, job_description, similarity_scores
                )
                
                results.append({
                    'filename': filename,
                    'scores': {
                        'sbert': round(sbert_score, 4),
                        'glove': round(glove_score, 4),
                        'doc2vec': round(doc2vec_score, 4),
                        'average': round(avg_score, 4)
                    },
                    'matched_skills': gemini_insights['matched_skills'],
                    'recommendations': gemini_insights['recommendations'],
                    'assessment': gemini_insights['assessment']
                })
                
                # Cleanup
                os.remove(filepath)
                print(f"🗑️ Temporary file removed")
        
        # Sort by average score
        results.sort(key=lambda x: x['scores']['average'], reverse=True)
        
        print("\n" + "="*50)
        print(f"✅ Analysis Complete! Processed {len(results)} resumes")
        print("="*50 + "\n")
        
        return jsonify({
            'success': True,
            'results': results
        })
        
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({
        'status': 'healthy',
        'models_loaded': True,
        'sbert': 'loaded',
        'glove': 'loaded' if glove_model is not None else 'will load on first use'
    })

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))

    print("\n" + "="*50)
    print("🎉 AI Resume Matcher Server")
    print("="*50)
    print(f"🚀 Server starting on port {port}")
    print("🌐 Running on Render / Docker")
    print("="*50 + "\n")

    app.run(host="0.0.0.0", port=port)
